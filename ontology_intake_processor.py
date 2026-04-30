"""
Ontology Intake Processor
=========================

Reads a folder (or uploaded files) containing:
  - CSV files       → tabular data (sampled for the LLM prompt)
  - PDF files       → data-model diagrams / ERDs / documentation
  - DOCX files      → data dictionaries / glossaries

Produces a single combined context string suitable for the existing
``build_user_prompt`` / ``call_azure_openai`` pipeline in
``llm_ontology_generator.py``.

Dependencies:
    pip install pdfplumber python-docx
"""

from __future__ import annotations

import csv
import os
from io import BytesIO
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# PDF text extraction (pdfplumber – layout-aware, table-friendly)
# ---------------------------------------------------------------------------

def extract_pdf_text(path_or_bytes: str | bytes | BytesIO) -> str:
    """Return the full text content of a PDF, page by page."""
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError(
            "pdfplumber is required for PDF extraction.  "
            "Install with: pip install pdfplumber"
        ) from e

    pages: list[str] = []
    if isinstance(path_or_bytes, (bytes, BytesIO)):
        buf = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
        with pdfplumber.open(buf) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- PDF page {i} ---\n{text.strip()}")
    else:
        with pdfplumber.open(path_or_bytes) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- PDF page {i} ---\n{text.strip()}")
    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# DOCX text extraction (python-docx)
# ---------------------------------------------------------------------------

def extract_docx_text(path_or_bytes: str | bytes | BytesIO) -> str:
    """Return the full text of a Word document (paragraphs + tables)."""
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction.  "
            "Install with: pip install python-docx"
        ) from e

    if isinstance(path_or_bytes, (bytes, BytesIO)):
        buf = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
        document = docx.Document(buf)
    else:
        document = docx.Document(path_or_bytes)

    parts: list[str] = []

    # Paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                prefix = "#" * (int(level) if level.isdigit() else 1)
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

    # Tables (data dictionaries are often in tables)
    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            parts.append("\n[TABLE]")
            for row in rows:
                parts.append(" | ".join(row))
            parts.append("[/TABLE]")

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# CSV sampling (reuses the same approach as llm_ontology_generator)
# ---------------------------------------------------------------------------

def sample_csv_text(path_or_bytes: str | bytes | BytesIO,
                    filename: str = "data.csv",
                    max_rows: int = 8) -> str:
    """Return a small textual sample of a CSV — header + first N rows."""
    if isinstance(path_or_bytes, (bytes, BytesIO)):
        buf = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
        text = buf.read().decode("utf-8", errors="replace")
        lines = text.splitlines()
    else:
        p = Path(path_or_bytes)
        if not p.exists():
            return f"(file not found: {path_or_bytes})"
        lines = p.read_text(encoding="utf-8", errors="replace").splitlines()

    rows: list[list[str]] = []
    reader = csv.reader(lines)
    for i, row in enumerate(reader):
        rows.append(row)
        if i >= max_rows:
            break

    if not rows:
        return f"(empty CSV: {filename})"

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header, *body = rows
    out = [f"=== CSV: {filename} ===",
           "columns: " + ", ".join(header),
           "sample rows:"]
    for r in body:
        out.append(" | ".join(r))
    return "\n".join(out)


# ---------------------------------------------------------------------------
# Folder scanner
# ---------------------------------------------------------------------------

def scan_intake_folder(folder: str | Path) -> dict[str, list[Path]]:
    """Categorize files in the intake folder by type."""
    folder = Path(folder)
    result: dict[str, list[Path]] = {"csv": [], "pdf": [], "docx": [], "other": []}
    if not folder.is_dir():
        return result
    for f in sorted(folder.iterdir()):
        if f.is_file():
            ext = f.suffix.lower()
            if ext == ".csv":
                result["csv"].append(f)
            elif ext == ".pdf":
                result["pdf"].append(f)
            elif ext in (".docx", ".doc"):
                result["docx"].append(f)
            elif ext not in (".md",):  # skip markdown files
                result["other"].append(f)
    return result


# ---------------------------------------------------------------------------
# Build combined context from a folder
# ---------------------------------------------------------------------------

def build_intake_context(folder: str | Path) -> str:
    """Read all CSVs, PDFs, and DOCX files from the intake folder and
    return a single combined context string for the LLM prompt."""
    files = scan_intake_folder(folder)
    sections: list[str] = []

    # CSVs — tabular data samples
    for csv_path in files["csv"]:
        sections.append(sample_csv_text(str(csv_path), filename=csv_path.name))

    # PDFs — data model / ERD documentation
    for pdf_path in files["pdf"]:
        try:
            text = extract_pdf_text(str(pdf_path))
            if text.strip():
                sections.append(f"=== PDF: {pdf_path.name} ===\n{text}")
        except Exception as e:
            sections.append(f"=== PDF: {pdf_path.name} === (extraction error: {e})")

    # DOCX — data dictionary / glossary
    for docx_path in files["docx"]:
        try:
            text = extract_docx_text(str(docx_path))
            if text.strip():
                sections.append(f"=== DOCX: {docx_path.name} ===\n{text}")
        except Exception as e:
            sections.append(f"=== DOCX: {docx_path.name} === (extraction error: {e})")

    return "\n\n" + "\n\n".join(sections) if sections else ""


# ---------------------------------------------------------------------------
# Build combined context from uploaded files (Streamlit UploadedFile objects)
# ---------------------------------------------------------------------------

def build_uploaded_context(uploaded_files: list[Any]) -> str:
    """Process a list of Streamlit UploadedFile objects and return
    a combined context string for the LLM prompt."""
    sections: list[str] = []

    for uf in uploaded_files:
        name = uf.name
        ext = Path(name).suffix.lower()
        data = uf.getvalue()

        if ext == ".csv":
            sections.append(sample_csv_text(data, filename=name))
        elif ext == ".pdf":
            try:
                text = extract_pdf_text(data)
                if text.strip():
                    sections.append(f"=== PDF: {name} ===\n{text}")
            except Exception as e:
                sections.append(f"=== PDF: {name} === (extraction error: {e})")
        elif ext in (".docx", ".doc"):
            try:
                text = extract_docx_text(data)
                if text.strip():
                    sections.append(f"=== DOCX: {name} ===\n{text}")
            except Exception as e:
                sections.append(f"=== DOCX: {name} === (extraction error: {e})")
        else:
            # Try to read as plain text
            try:
                text = data.decode("utf-8", errors="replace")
                if text.strip():
                    sections.append(f"=== FILE: {name} ===\n{text[:3000]}")
            except Exception:
                sections.append(f"=== FILE: {name} === (unsupported format)")

    return "\n\n" + "\n\n".join(sections) if sections else ""


# ---------------------------------------------------------------------------
# Enhanced prompt builder for intake scenarios
# ---------------------------------------------------------------------------

INTAKE_SYSTEM_PROMPT_ADDENDUM = """

ADDITIONAL CONTEXT — DATA-MODEL-DRIVEN ONTOLOGY:
You are also provided with:
  - Sample rows from CSV data files (the actual instance data).
  - A data-model document (PDF) describing entities, relationships, and ERD.
  - A data-dictionary document (DOCX) with column definitions, data types,
    constraints, and business rules.

Use ALL of these sources together to build the ontology:
  1. CSV column names and sample data → infer classes, data properties, and
     controlled vocabularies (SKOS-style concept schemes).
  2. PDF data model → capture entity relationships, cardinalities, and the
     conceptual hierarchy (foreign keys → object properties).
  3. DOCX data dictionary → enrich property comments with business definitions,
     constraints, and validation rules.
  4. Reconcile naming across sources — the data dictionary is authoritative
     for business names; use PascalCase versions as class/property names.
  5. Mark foreign-key relationships as object properties with correct
     domain/range.
  6. Create a class for each table/entity, plus meaningful superclasses to
     ensure a deep hierarchy (aim for 4-6 levels).
"""
